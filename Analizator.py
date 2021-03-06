import sys
from PyQt5 import QtCore, QtGui, QtWidgets, Qt, uic
from PyQt5.QtWidgets import *
import traceback
import pandas as pd
import time
import requests
from docx import Document
import matplotlib.pyplot as plt
from matplotlib.backends.backend_pdf import PdfPages

def log_uncaught_exceptions(ex_cls, ex, tb): #error catcher
    text = '{}: {}:\n'.format(ex_cls.__name__, ex)
    text += ''.join(traceback.format_tb(tb))
    QtWidgets.QMessageBox.critical(None, 'Error', text)
    sys.exit()
sys.excepthook = log_uncaught_exceptions

lines = []
diction = {}

Form, _ = uic.loadUiType('main.ui')
filter, _ = uic.loadUiType('filter.ui')


class win_filter(QMainWindow, filter):
    global chose_err
    chose_err = []
    def __init__(self, parent=None):
        super().__init__(parent, QtCore.Qt.Window)
        self.parent = parent
        self.setupUi(self)

        self.upload_docx.clicked.connect(lambda: self.upload_to_docx())
        self.show_diagramm.clicked.connect(lambda: self.show_diagr())
        dict_for_err = []
        for i in diction:
            if str(diction[i][8]) not in dict_for_err:
                self.listWidget.addItem(str(diction[i][8]))
                dict_for_err.append(str(diction[i][8]))

    def collect_from_filters(self):
        chose_err.clear()
        global year_s, year_f, status, OS__1, OS__2, OS__3, exploit, check_box_state
        year_s = str(self.date_start.text()).split(".")[0]
        year_f = str(self.date_finish.text()).split(".")[0]
        status = self.status_yazvim.currentText()
        OS__1 =self.OS_1.currentText()
        OS__2 = self.OS_2.currentText()
        OS__3 = self.OS_3.currentText()
        exploit = self.explouit.currentText()
        check_box_state = self.without_errors.checkState()
        if int(check_box_state) != 2:
            for i in range(int(self.chosen_errors.count())):
                chose_err.append(self.chosen_errors.item(i).text())

    def filtering(self):
        global collect
        collect = {}
        self.collect_from_filters()
        for i in diction:
            checker = {"date": False, "status": False,
                       "OS__1": False, "OS__2": False, "OS__3": False,
                       "exploit": False, "errors": False}
            pr_year = str(diction[i][3].split(".")[2])

            if OS__2 == "Без фильтра":
                checker.pop("OS__2")

            if OS__3 == "Без фильтра":
                checker.pop("OS__3")

            def check_date():
                if year_s <= pr_year <= year_f:
                    checker["date"] = True
                else:
                    pass

            def check_status():
                stat = {"Уязвимость устранена":"Уязвимость устранена", "Уязвимость актуальна":"Информация об устранении отсутствует"}
                print(stat[status])
                print(diction[i][7])
                if stat[status] == diction[i][7]:
                    checker["status"] = True

            def check_OS():
                if str(diction[i][1]) != "nan":
                    if OS__1 in diction[i][1]:
                        checker["OS__1"] = True
                    if OS__2 in diction[i][1]:
                        checker["OS__2"] = True
                    if OS__3 in diction[i][1]:
                        checker["OS__3"] = True

            def check_exploit():
                expl = {"Существует":"Существует", "Не существует":"Данные уточняются", "Существует":"Существует в открытом доступе"}
                print(expl[exploit])
                print(diction[i][6])
                if expl[exploit] == diction[i][6]:
                    checker["exploit"] = True

            def check_errors():
                if diction[i][8] in chose_err:
                    checker["errors"] = True
                else:
                    checker["errors"] = False

            check_date()

            if status == "Без фильтра":
                checker.pop("status")
            else:
                check_status()

            check_OS()

            if exploit == "Без фильтра":
                checker.pop("exploit")
            else:
                check_exploit()

            if int(check_box_state) != 2:
                check_errors()
            else:
                checker.pop("errors")

            pass_flag = False
            for g in checker:
                if checker[g] == False:
                        pass_flag = True
                        break

            try:
                if diction[i]["status"] != status:
                    pass_flag = True
            except Exception:
               None

            try:
                if diction[i]["exploit"] != status:
                    pass_flag = True
            except Exception:
                None

            if pass_flag == False:
                print("True")
                collect[i] = diction[i]


    def upload_to_docx(self):
        try:
            self.filtering()
            doc = Document()
            doc.add_paragraph('Таблица вывода,\nK-Критический уровень опасности,\nH-Высокий уровень опасности,\nM-Средний уровень опасности,\nL-Низкий уровень опасности')
            table = doc.add_table(rows=len(collect)+1, cols=6)
            table.cell(0, 0).text = 'Дата'
            table.cell(0, 1).text = 'Статус\nуязвимости'
            table.cell(0, 2).text = 'ОС'
            table.cell(0, 3).text = 'Наличие\nЭксплойта'
            table.cell(0, 4).text = 'Ошибки'
            table.cell(0, 5).text = 'Ур.\n опасности'
            codes_to_dang = {"Критический уровень опасности ":"K", "Высокий уровень опасности ":"H", "Средний уровень опасности ":"M", "Низкий уровень опасности ":"L"}
            counter = 1
            for i in collect:
                dang =codes_to_dang[collect[i][4].split("(")[0]]
                table.cell(counter, 0).text = collect[i][3]
                table.cell(counter, 1).text = collect[i][7]
                table.cell(counter, 2).text = collect[i][1]
                table.cell(counter, 3).text = collect[i][6]
                table.cell(counter, 4).text = collect[i][8]
                table.cell(counter, 5).text = dang
                counter += 1
            doc.save('report.docx')

            self.message = QMessageBox.information(self,'Vse' ,'Задача завершена ')

        except IOError:
            self.message = QMessageBox.critical(self,'error' ,'Внимание, у вас открыт файл,\nзайкройте его и повторите операцию')

        """
        #name[0], date[3], status[7], oc[1], oc, oc exploit[6], errors[8]   ('Adobe Reader, Adobe Reader Document Cloud, Adobe Acrobat Document Cloud, Adobe Acrobat', 'Mac OS (X), Windows (.)', 'Уязвимость кода', '11.05.2016', 'Критический уровень опасности (базовая оценка CVSS 2.0 составляет 10)', 'Подтверждена производителем', 'Данные уточняются', 'Уязвимость устранена', 'CWE-416')
        """

    def show_diagr(self):
        self.filtering()
        index = []
        values_years = []
        for i in range(int(year_s), int(year_f)+1):
            index.append(i)
            values_years.append(code_to_date_obn[str(i)])
        print(index)
        print(values_years)
        plt.bar(index, values_years)
        plt.savefig('graph.png', bbox_inches='tight', dpi= 150)
        plt.show()

        ######   codes_to_yr_opas, codes_to_ystr_yaz, codes_to_status_yaz, codes_to_class_yazvim, code_to_date_ystr, code_to_date_obn
        #codes_to_yr_opas = {"Критический уровень опасности ":krit_yaz, "Высокий уровень опасности ":high_yaz, "Средний уровень опасности ":med_yaz, "Низкий уровень опасности ":low_yaz}

class main_ui(QMainWindow, Form):
    def __init__(self):
        super(main_ui, self).__init__()
        self.setupUi(self)
        self.initUI()


    def initUI(self):
        self.build_diagr.clicked.connect(lambda: self.open_filt())
        self.load_file.clicked.connect(lambda: self.load_file_fun())
        self.upload_new_base.clicked.connect(lambda: self.upload())
        self.message.clicked.connect(lambda: QMessageBox.information(self, 'Баг-репорт', "Не баг, а фича!"))

    def upload(self):
        self.message = QMessageBox.information(self, 'Обновление', "Начало обновления базы угроз\nНе закрывайте приложение")
        url = 'https://bdu.fstec.ru/files/documents/vullist.xlsx'
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) '
                          'Chrome/91.0.4472.114 Safari/537.36',
            'Referrer': 'https://bdu.fstec.ru/vul'}
        result = requests.get(url, headers=headers)
        output = open('vullist.xlsx', 'wb')
        output.write(result.content)
        output.close()
        self.message = QMessageBox.information(self, 'Обновление', "Обновление базы успешно")

    def open_filt(self):
        ui2 = win_filter(self)
        ui2.show()

    def load_file_fun(self):
        chose_priloj = self.chose_pril.text()
        col = [4, 7, 8, 9, 12, 14, 15, 16, 21]

        file, __ = QFileDialog.getOpenFileName(self, 'Open file')
        self.message = QMessageBox.information(self, 'Жди' ,'Подгружаем файл, ожидайте')
        try:
            if file != "":
                start_time = time.time()

                top_players = pd.read_excel(file, usecols=col)

                counter = -1
                for name in top_players['Unnamed: 4'].tolist():
                    counter += 1
                    if "Adobe Reader" in str(name):
                        lines.append(counter)
                        diction[counter] = str(chose_priloj)
                for i in lines:
                    programm = top_players['Unnamed: 4'].tolist()[i]
                    os = top_players['Unnamed: 7'].tolist()[i]
                    class_yazvim = top_players['Unnamed: 8'].tolist()[i]
                    date =  top_players['Unnamed: 9'].tolist()[i]
                    yr_opas =  top_players['Unnamed: 12'].tolist()[i]
                    status_yaz =  top_players['Unnamed: 14'].tolist()[i]
                    nal_exploit =  top_players['Unnamed: 15'].tolist()[i]
                    ystr_yaz =  top_players['Unnamed: 16'].tolist()[i]
                    type_error =  top_players['Unnamed: 21'].tolist()[i]
                    diction[i] = (programm, os , class_yazvim, date, yr_opas, status_yaz, nal_exploit, ystr_yaz, type_error)

                krit_yaz = 0
                high_yaz = 0
                med_yaz = 0
                low_yaz = 0

                krit_yaz_ystr = 0
                high_yaz_ystr = 0
                med_yaz_ystr = 0
                low_yaz_ystr = 0
                none1 = 0
                none2 = 0
                none3 = 0
                none4 = 0

                krit_yaz_podtv = 0
                high_yaz_podtv = 0
                med_yaz_podtv = 0
                low_yaz_podtv = 0
                krit_yaz_potenc = 0
                high_yaz_potenc = 0
                med_yaz_potenc = 0
                low_yaz_potenc = 0

                class_yazvim_arch = 0
                class_yazvim_code = 0
                class_yazvim_many = 0
                class_yazvim_none = 0


                y1999A = 0
                y1999Y = 0
                y2000A = 0
                y2000Y = 0
                y2001A = 0
                y2001Y = 0
                y2002A = 0
                y2002Y = 0
                y2003A = 0
                y2003Y = 0
                y2004A = 0
                y2004Y = 0
                y2005A = 0
                y2005Y = 0
                y2006A = 0
                y2006Y = 0
                y2007A = 0
                y2007Y = 0
                y2008A = 0
                y2008Y = 0
                y2009A = 0
                y2009Y = 0
                y2010A = 0
                y2010Y = 0
                y2011A = 0
                y2011Y = 0
                y2012A = 0
                y2012Y = 0
                y2013A = 0
                y2013Y = 0
                y2014A = 0
                y2014Y = 0
                y2015A = 0
                y2015Y = 0
                y2016A = 0
                y2016Y = 0
                y2017A = 0
                y2017Y = 0
                y2018A = 0
                y2018Y = 0
                y2019A = 0
                y2019Y = 0
                y2020A = 0
                y2020Y = 0
                y2021A = 0
                y2021Y = 0
                global codes_to_yr_opas, codes_to_ystr_yaz, codes_to_status_yaz, codes_to_class_yazvim, code_to_date_ystr, code_to_date_obn

                codes_to_yr_opas = {"Критический уровень опасности ":krit_yaz, "Высокий уровень опасности ":high_yaz, "Средний уровень опасности ":med_yaz, "Низкий уровень опасности ":low_yaz}
                codes_to_ystr_yaz = {"Уязвимость устранена":{"Критический уровень опасности ":krit_yaz_ystr, "Высокий уровень опасности ":high_yaz_ystr, "Средний уровень опасности ":med_yaz_ystr, "Низкий уровень опасности ":low_yaz_ystr},
                                     'Информация об устранении отсутствует':{"Критический уровень опасности ":none1, "Высокий уровень опасности ":none2, "Средний уровень опасности ":none3, "Низкий уровень опасности ":none4}}
                codes_to_status_yaz = {"Подтверждена":{"Критический уровень опасности ":krit_yaz_podtv, "Высокий уровень опасности ":high_yaz_podtv, "Средний уровень опасности ":med_yaz_podtv, "Низкий уровень опасности ":low_yaz_podtv},
                                       "Потенциальная":{"Критический уровень опасности ":krit_yaz_potenc, "Высокий уровень опасности ":high_yaz_potenc, "Средний уровень опасности ":med_yaz_potenc, "Низкий уровень опасности ":low_yaz_potenc}}
                codes_to_class_yazvim = {"Уязвимость архитектуры":class_yazvim_arch, "Уязвимость кода":class_yazvim_code, "Уязвимость многофакторная":class_yazvim_many, "nan":class_yazvim_none}
                code_to_date_ystr = {"1999":{"Уязвимость устранена":y1999Y, 'Информация об устранении отсутствует':none4}, "2000":{"Уязвимость устранена":y2000Y, 'Информация об устранении отсутствует':none4}, "2001":{"Уязвимость устранена":y2001Y, 'Информация об устранении отсутствует':none4},
                                     "2002":{"Уязвимость устранена":y2002Y, 'Информация об устранении отсутствует':none4}, "2003":{"Уязвимость устранена":y2003Y, 'Информация об устранении отсутствует':none4}, "2004":{"Уязвимость устранена":y2004Y, 'Информация об устранении отсутствует':none4},
                                     "2005":{"Уязвимость устранена":y2005Y, 'Информация об устранении отсутствует':none4}, "2006":{"Уязвимость устранена":y2006Y, 'Информация об устранении отсутствует':none4}, "2007":{"Уязвимость устранена":y2007Y, 'Информация об устранении отсутствует':none4},
                                     "2008":{"Уязвимость устранена":y2008Y, 'Информация об устранении отсутствует':none4}, "2009":{"Уязвимость устранена":y2009Y, 'Информация об устранении отсутствует':none4}, "2010":{"Уязвимость устранена":y2010Y, 'Информация об устранении отсутствует':none4},
                                     "2011":{"Уязвимость устранена":y2011Y, 'Информация об устранении отсутствует':none4}, "2012":{"Уязвимость устранена":y2012Y, 'Информация об устранении отсутствует':none4}, "2013":{"Уязвимость устранена":y2013Y, 'Информация об устранении отсутствует':none4},
                                     "2014":{"Уязвимость устранена":y2014Y, 'Информация об устранении отсутствует':none4}, "2015":{"Уязвимость устранена":y2015Y, 'Информация об устранении отсутствует':none4}, "2016":{"Уязвимость устранена":y2016Y, 'Информация об устранении отсутствует':none4},
                                     "2017":{"Уязвимость устранена":y2017Y, 'Информация об устранении отсутствует':none4}, "2018":{"Уязвимость устранена":y2018Y, 'Информация об устранении отсутствует':none4}, "2019":{"Уязвимость устранена":y2019Y, 'Информация об устранении отсутствует':none4},
                                     "2020":{"Уязвимость устранена":y2020Y, 'Информация об устранении отсутствует':none4}, "2021":{"Уязвимость устранена":y2021Y, 'Информация об устранении отсутствует':none4}}
                code_to_date_obn = {"1999":y1999A, "2000":y2000A, "2001":y2001A, "2002":y2002A, "2003":y2003A, "2004":y2004A, "2005":y2005A, "2006":y2006A, "2007":y2007A, "2008":y2008A, "2009":y2009A, "2010":y2010A,
                                    "2011":y2011A, "2012":y2012A, "2013":y2013A, "2014":y2014A, "2015":y2015A,"2016":y2016A,"2017":y2017A, "2018":y2018A,"2019":y2019A, "2020":y2020A,"2021":y2021A}

                for i in diction:
                    codes_to_yr_opas[diction[i][4].split("(")[0]] += 1
                    codes_to_ystr_yaz[diction[i][7].split("(")[0]][diction[i][4].split("(")[0]] += 1
                    codes_to_status_yaz[diction[i][5].split(" ")[0]][diction[i][4].split("(")[0]] += 1
                    codes_to_class_yazvim[str(diction[i][2])] += 1
                    code_to_date_ystr[str(diction[i][3].split(".")[2])][str(diction[i][7].split("(")[0])] += 1
                    code_to_date_obn[str(diction[i][3].split(".")[2])] += 1

                self.krit_obn_labl.setText(str(codes_to_yr_opas["Критический уровень опасности "]))
                self.hig_obn_labl.setText(str(codes_to_yr_opas["Высокий уровень опасности "]))
                self.med_obn_labl.setText(str(codes_to_yr_opas["Средний уровень опасности "]))
                self.low_obn_labl.setText(str(codes_to_yr_opas["Низкий уровень опасности "]))

                self.krit_ystr_labl.setText(str(codes_to_ystr_yaz["Уязвимость устранена"]["Критический уровень опасности "]))
                self.hig_ystr_labl.setText(str(codes_to_ystr_yaz["Уязвимость устранена"]["Высокий уровень опасности "]))
                self.med_ystr_labl.setText(str(codes_to_ystr_yaz["Уязвимость устранена"]["Средний уровень опасности "]))
                self.low_ystr_labl.setText(str(codes_to_ystr_yaz["Уязвимость устранена"]["Низкий уровень опасности "]))

                self.krit_podtvpot_labl.setText(str(codes_to_status_yaz["Подтверждена"]["Критический уровень опасности "])+"/"+str(codes_to_status_yaz["Потенциальная"]["Критический уровень опасности "]))
                self.hig_podtvpot_labl.setText(str(codes_to_status_yaz["Подтверждена"]["Высокий уровень опасности "])+"/"+str(codes_to_status_yaz["Потенциальная"]["Высокий уровень опасности "]))
                self.mid_podtvpot_labl.setText(str(codes_to_status_yaz["Подтверждена"]["Средний уровень опасности "])+"/"+str(codes_to_status_yaz["Потенциальная"]["Средний уровень опасности "]))
                self.low_podtvpot_labl.setText(str(codes_to_status_yaz["Подтверждена"]["Низкий уровень опасности "])+"/"+str(codes_to_status_yaz["Потенциальная"]["Низкий уровень опасности "]))

                self.yiaz_ache.setText(str(codes_to_class_yazvim["Уязвимость архитектуры"]))
                self.yiaz_code.setText(str(codes_to_class_yazvim["Уязвимость кода"]))
                self.yiaz_manyfact.setText(str(codes_to_class_yazvim["Уязвимость многофакторная"]))
                self.yiaz_none.setText(str(codes_to_class_yazvim["nan"]))


                self.obn_1999.setText(str(code_to_date_obn["1999"]))
                self.obn_2000.setText(str(code_to_date_obn["2000"]))
                self.obn_2001.setText(str(code_to_date_obn["2001"]))
                self.obn_2002.setText(str(code_to_date_obn["2002"]))
                self.obn_2003.setText(str(code_to_date_obn["2003"]))
                self.obn_2004.setText(str(code_to_date_obn["2004"]))
                self.obn_2005.setText(str(code_to_date_obn["2005"]))
                self.obn_2006.setText(str(code_to_date_obn["2006"]))
                self.obn_2007.setText(str(code_to_date_obn["2007"]))
                self.obn_2008.setText(str(code_to_date_obn["2008"]))
                self.obn_2009.setText(str(code_to_date_obn["2009"]))
                self.obn_2010.setText(str(code_to_date_obn["2010"]))
                self.obn_2011.setText(str(code_to_date_obn["2011"]))
                self.obn_2012.setText(str(code_to_date_obn["2012"]))
                self.obn_2013.setText(str(code_to_date_obn["2013"]))
                self.obn_2014.setText(str(code_to_date_obn["2014"]))
                self.obn_2015.setText(str(code_to_date_obn["2015"]))
                self.obn_2016.setText(str(code_to_date_obn["2016"]))
                self.obn_2017.setText(str(code_to_date_obn["2017"]))
                self.obn_2018.setText(str(code_to_date_obn["2018"]))
                self.obn_2019.setText(str(code_to_date_obn["2019"]))
                self.obn_2020.setText(str(code_to_date_obn["2020"]))
                self.obn_2021.setText(str(code_to_date_obn["2021"]))
                #
                self.ystr_1999.setText(str(code_to_date_ystr["1999"]["Уязвимость устранена"]))
                self.ystr_2000.setText(str(code_to_date_ystr["2000"]["Уязвимость устранена"]))
                self.ystr_2001.setText(str(code_to_date_ystr["2001"]["Уязвимость устранена"]))
                self.ystr_2002.setText(str(code_to_date_ystr["2002"]["Уязвимость устранена"]))
                self.ystr_2003.setText(str(code_to_date_ystr["2003"]["Уязвимость устранена"]))
                self.ystr_2004.setText(str(code_to_date_ystr["2004"]["Уязвимость устранена"]))
                self.ystr_2005.setText(str(code_to_date_ystr["2005"]["Уязвимость устранена"]))
                self.ystr_2006.setText(str(code_to_date_ystr["2006"]["Уязвимость устранена"]))
                self.ystr_2007.setText(str(code_to_date_ystr["2007"]["Уязвимость устранена"]))
                self.ystr_2008.setText(str(code_to_date_ystr["2008"]["Уязвимость устранена"]))
                self.ystr_2009.setText(str(code_to_date_ystr["2009"]["Уязвимость устранена"]))
                self.ystr_2010.setText(str(code_to_date_ystr["2010"]["Уязвимость устранена"]))
                self.ystr_2011.setText(str(code_to_date_ystr["2011"]["Уязвимость устранена"]))
                self.ystr_2012.setText(str(code_to_date_ystr["2012"]["Уязвимость устранена"]))
                self.ystr_2013.setText(str(code_to_date_ystr["2013"]["Уязвимость устранена"]))
                self.ystr_2014.setText(str(code_to_date_ystr["2014"]["Уязвимость устранена"]))
                self.ystr_2015.setText(str(code_to_date_ystr["2015"]["Уязвимость устранена"]))
                self.ystr_2016.setText(str(code_to_date_ystr["2016"]["Уязвимость устранена"]))
                self.ystr_2017.setText(str(code_to_date_ystr["2017"]["Уязвимость устранена"]))
                self.ystr_2018.setText(str(code_to_date_ystr["2018"]["Уязвимость устранена"]))
                self.ystr_2019.setText(str(code_to_date_ystr["2019"]["Уязвимость устранена"]))
                self.ystr_2020.setText(str(code_to_date_ystr["2020"]["Уязвимость устранена"]))
                self.ystr_2021.setText(str(code_to_date_ystr["2021"]["Уязвимость устранена"]))

                chtenie = str((time.time() - start_time))
                self.message = QMessageBox.information(self, 'На чтение затрачено', 'На чтение затрачено '+ chtenie + ' секунд')
        except Exception:
            self.message = QMessageBox.warning(self, 'Ошибка', 'Ошибка')
if __name__ == "__main__":
    app = QtWidgets.QApplication(sys.argv)
    ui1 = main_ui()
    ui1.show()
    sys.exit(app.exec_())